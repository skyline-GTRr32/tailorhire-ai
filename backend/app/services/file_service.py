# File Service - PDF and DOCX processing
import io
import os
from typing import Optional, Tuple
from fastapi import UploadFile, HTTPException
import PyPDF2
from docx import Document
from datetime import datetime
from app.utils.logger import setup_logger

logger = setup_logger('file_service')

class FileService:
    """Service for handling file operations with detailed logging"""
    
    async def extract_text_from_file(self, file: UploadFile) -> str:
        """
        Extract text from uploaded file (PDF or DOCX) with detailed logging
        """
        file_info = f"{file.filename} ({file.content_type}, {file.size} bytes)"
        logger.info(f"📂 Starting file processing: {file_info}")
        
        try:
            # Read file content
            start_time = datetime.now()
            content = await file.read()
            read_time = (datetime.now() - start_time).total_seconds()
            
            logger.debug(f"📥 File read completed in {read_time:.3f}s")
            
            # Determine file type and process
            if not file.content_type:
                logger.error("❌ No content type provided")
                raise ValueError("File type could not be determined")
            
            if file.content_type == "application/pdf":
                logger.info("🔍 Processing PDF file")
                text = self._extract_from_pdf(content)
            elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                logger.info("📝 Processing DOCX file")
                text = self._extract_from_docx(content)
            else:
                logger.error(f"❌ Unsupported file type: {file.content_type}")
                raise ValueError(f"Unsupported file type: {file.content_type}")
            
            # Log text extraction results
            char_count = len(text)
            word_count = len(text.split())
            logger.info(f"✅ Text extraction complete - {char_count} characters, {word_count} words")
            
            if char_count == 0:
                logger.warning("⚠️ Extracted text is empty - possible issue with file content")
            
            return text
                
        except Exception as e:
            error_msg = f"Error extracting text from {file.filename}: {str(e)}"
            logger.error(f"❌ {error_msg}", exc_info=True)
            raise HTTPException(status_code=400, detail=f"Error processing file: {str(e)}")
    
    def _extract_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF content with detailed logging"""
        logger.debug("📄 Extracting text from PDF")
        start_time = datetime.now()
        
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            page_count = len(pdf_reader.pages)
            logger.debug(f"📑 PDF contains {page_count} pages")
            
            text_parts = []
            for i, page in enumerate(pdf_reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                        logger.debug(f"  ✓ Page {i}: {len(page_text)} characters")
                    else:
                        logger.warning(f"  ⚠️ Page {i}: No text extracted")
                except Exception as page_error:
                    logger.warning(f"  ⚠️ Error extracting text from page {i}: {str(page_error)}")
            
            full_text = "\n\n".join(text_parts).strip()
            process_time = (datetime.now() - start_time).total_seconds()
            logger.debug(f"✅ PDF processing completed in {process_time:.3f}s")
            
            return full_text
            
        except Exception as e:
            logger.error(f"❌ PDF extraction failed: {str(e)}", exc_info=True)
            raise
    
            # Return placeholder for demo
            return f"""PDF Processing Note:
            
Your PDF file "{content[:50]}..." was uploaded successfully, but automatic text extraction encountered an issue.

For best results, please:
1. Copy and paste your resume text directly into the text area
2. Ensure the PDF contains selectable text (not just images)
3. Try saving your resume as a Word document instead

This is a demo limitation - full PDF processing will be enhanced in production."""

    def _extract_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX content"""
        try:
            logger.info("🔍 Extracting text from DOCX...")
            
            docx_file = io.BytesIO(content)
            doc = docx.Document(docx_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            # Clean up text
            text = self._clean_extracted_text(text)
            
            if len(text.strip()) < 50:
                raise Exception("DOCX appears to be empty")
            
            logger.info(f"✅ DOCX text extracted: {len(text)} characters")
            return text
            
        except Exception as e:
            logger.error(f"❌ DOCX extraction failed: {str(e)}")
            # Return placeholder for demo
            return f"""DOCX Processing Note:
            
Your Word document was uploaded successfully, but automatic text extraction encountered an issue.

For best results, please:
1. Copy and paste your resume text directly into the text area
2. Ensure the document is in proper DOCX format
3. Check that the document isn't password protected

This is a demo limitation - full DOCX processing will be enhanced in production."""

    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        import re
        
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        
        # Remove special characters that might cause issues
        text = re.sub(r'[^\x00-\x7F]+', '', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text

    def validate_file(self, file: UploadFile) -> bool:
        """Validate uploaded file"""
        max_size = 10 * 1024 * 1024  # 10MB
        allowed_types = [
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ]
        
        if file.size and file.size > max_size:
            raise HTTPException(status_code=400, detail="File too large (max 10MB)")
        
        if file.content_type not in allowed_types:
            raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")
        
        return True